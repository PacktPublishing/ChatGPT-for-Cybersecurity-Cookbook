import openai
from openai import OpenAI # Updated import statement
import os
from docx import Document
import threading
import time
from datetime import datetime
from tqdm import tqdm

# Set up the OpenAI API
openai.api_key = os.getenv("OPENAI_API_KEY")

current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
assessment_name = f"Risk_Assessment_Plan_{current_datetime}"

# Risk Assessment Outline
risk_assessment_outline = [
    "Define Business Objectives",
    "Asset Discovery/Identification",
    "System Characterization/Classification",
    "Network Diagrams and Data Flow Review",
    "Risk Pre-Screening",
    "Security Policy & Procedures Review",
    "Cybersecurity Standards Selection and Gap Assessment/Audit",
    "Vulnerability Assessment",
    "Threat Assessment",
    "Attack Vector Assessment",
    "Risk Scenario Creation (using the Mitre ATT&CK Framework)",
    "Validate Findings with Penetration Testing/Red Teaming",
    "Risk Analysis (Aggregate Findings & Calculate Risk Scores)",
    "Prioritize Risks",
    "Assign Mitigation Methods and Tasks",
    "Create Risk Report",
]

# Function to generate a section content using the OpenAI API
def generate_section_content(section: str) -> str:
    # Define the conversation messages
    messages = [
        {"role": "system", "content": 'You are a cybersecurity professional specializing in governance, risk, and compliance (GRC) with more than 25 years of experience.'},
        {"role": "user", "content": f'You are currently writing a cyber risk assessment policy. Write the narrative, context, and details for the following section (and only this section): {section}. Use as much detail and explanation as possible. Do not write anything that should go in another section of the policy.'},
    ]

    # Call the OpenAI API
    client = OpenAI()   # New client initialization required for the updated API call
    response = client.chat.completions.create( # Updated API call
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        temperature=0.5,
    )

    # Return the generated text
    return response.choices[0].message.content.strip() # Updated API call

# Function to convert markdown text to a Word document
def markdown_to_docx(markdown_text: str, output_file: str):
    document = Document()

    # Iterate through the lines of the markdown text
    for line in markdown_text.split('\n'):
        # Add headings based on the markdown heading levels
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        # Add paragraphs for other text
        else:
            document.add_paragraph(line)

    # Save the Word document
    document.save(output_file)

# Function to display elapsed time while waiting for the API call
def display_elapsed_time():
    start_time = time.time()
    while not api_call_completed:
        elapsed_time = time.time() - start_time
        print(f"\rElapsed time: {elapsed_time:.2f} seconds", end="")
        time.sleep(1)

api_call_completed = False
elapsed_time_thread = threading.Thread(target=display_elapsed_time)
elapsed_time_thread.start()

# Generate the report using the OpenAI API
report = []
pbar = tqdm(total=len(risk_assessment_outline), desc="Generating sections")
for section in risk_assessment_outline:
    try:
        # Generate the section content
        content = generate_section_content(section)
        # Append the section content to the report
        report.append(f"## {section}\n{content}")
    except Exception as e:
        print(f"\nAn error occurred during the API call: {e}")
        api_call_completed = True
        exit()
    pbar.update(1)

api_call_completed = True
elapsed_time_thread.join()
pbar.close()

# Save the report as a Word document
docx_output_file = f"{assessment_name}_report.docx"

# Handle exceptions during the report generation
try:
    markdown_to_docx('\n'.join(report), docx_output_file)
    print("\nReport generated successfully!")
except Exception as e:
    print(f"\nAn error occurred during the report generation: {e}")
