import os
import openai
import docx
from markdown import markdown
from tqdm import tqdm

# get the OpenAI API key from environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

# prepare initial prompt
messages=[
    {
        "role": "system",
        "content": "You are a cybersecurity professional specializing in governance, risk, and compliance (GRC) with more than 25 years of experience."
    },
    {
        "role": "user",
        "content": "Write a detailed cybersecurity policy outline for my company, XYZ Corp., which is a credit union. Provide the outline only, with no context or narrative. Use markdown language to denote the proper headings, lists, formatting, etc."
    }
]

print("Generating policy outline...")
try:
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )
except Exception as e:
    print("An error occurred while connecting to the OpenAI API:", e)
    exit(1)

# get outline
outline = response['choices'][0]['message']['content'].strip()

print(outline + "\n")

# split outline into sections
sections = outline.split("\n\n")

# prepare Word document
doc = docx.Document()
html_text = ""

# for each section in the outline
for i, section in tqdm(enumerate(sections, start=1), total=len(sections), leave=False):
    print(f"\nGenerating details for section {i}...")

    # prepare prompt for detailed info
    messages=[
        {
            "role": "system",
            "content": "You are a cybersecurity professional specializing in governance, risk, and compliance (GRC) with more than 25 years of experience."
        },
        {
            "role": "user",
            "content": f"You are currently writing a cybersecurity policy. Write the narrative, context, and details for the following section (and only this section): {section}. Use as much detail and explanation as possible. Do not write anything that should go in another section of the policy."
        }
    ]
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=2048,
            n=1,
            stop=None,
            temperature=0.7,
        )
    except Exception as e:
        print("An error occurred while connecting to the OpenAI API:", e)
        exit(1)

    # get detailed info
    detailed_info = response['choices'][0]['message']['content'].strip()

    # convert markdown to Word formatting
    doc.add_paragraph(detailed_info)
    doc.add_paragraph("\n")  # add extra line break for readability

    # convert markdown to HTML and add to the html_text string
    html_text += markdown(detailed_info)

    # save Word document
    print("Saving sections...")
    doc.save("Cybersecurity_Policy.docx")

    # save HTML document
    with open("Cybersecurity_Policy.html", 'w') as f:
        f.write(html_text)

print("\nDone.")