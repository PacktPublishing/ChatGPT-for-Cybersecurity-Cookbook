import os
import openai
from openai import OpenAI # New import required for the updated API call
import docx
from markdown import markdown
from tqdm import tqdm

# get the OpenAI API key from environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

# prepare initial prompt
messages=[
    {
        "role": "system",
        "content": "You are a cybersecurity professional specializing in governance, risk, and compliance (GRC) with more than 25 years of experience."
    },
    {
        "role": "user",
        "content": "Write a detailed cybersecurity policy outline for my company, XYZ Corp., which is a water and wastewater facility. Provide the outline only, with no context or narrative. Use markdown language to denote the proper headings, lists, formatting, etc."
    }
]

print("Generating policy outline...")
try:
    client = OpenAI() # New client initialization required for the updated API call
    response = client.chat.completions.create( # Updated API call
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=0.5,
    )
except Exception as e:
    print("An error occurred while connecting to the OpenAI API:", e)
    exit(1)

# get outline
outline = response.choices[0].message.content.strip() # Updated API call

print(outline + "\n")

# split outline into sections
sections = outline.split("\n\n")

# prepare Word document
doc = docx.Document()
html_text = ""

# for each section in the outline
for i, section in tqdm(enumerate(sections, start=1), total=len(sections), leave=False):
    print(f"\nGenerating details for section {i}...")

    # prepare prompt for detailed info
    messages=[
        {
            "role": "system",
            "content": "You are a cybersecurity professional specializing in governance, risk, and compliance (GRC) with more than 25 years of experience."
        },
        {
            "role": "user",
            "content": f"You are currently writing a cybersecurity policy. Write the narrative, context, and details for the following section (and only this section): {section}. Use as much detail and explanation as possible. Do not write anything that should go in another section of the policy."
        }
    ]
    
    try:
        response = client.chat.completions.create( # Updated API call
            model="gpt-4-1106-preview",
            messages=messages,
            temperature=0.5,
        )
    except Exception as e:
        print("An error occurred while connecting to the OpenAI API:", e)
        exit(1)

    # get detailed info
    detailed_info = response.choices[0].message.content.strip() # Updated API call

    # convert markdown to Word formatting
    doc.add_paragraph(detailed_info)
    doc.add_paragraph("\n")  # add extra line break for readability

    # convert markdown to HTML and add to the html_text string
    html_text += markdown(detailed_info)

    # save Word document
    print("Saving sections...")
    doc.save("Cybersecurity_Policy.docx")

    # save HTML document
    with open("Cybersecurity_Policy.html", 'w') as f:
        f.write(html_text)

print("\nDone.")
