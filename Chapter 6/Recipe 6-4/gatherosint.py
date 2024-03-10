import openai
from openai import OpenAI
import os
from docx import Document
from datetime import datetime

client = OpenAI()  

# Set up the OpenAI API
openai.api_key = os.getenv("OPENAI_API_KEY")

def prompt_osint_analysis(job_desc):
    # Define the conversation messages
    messages = [
        {"role": "system", "content": 'You are a cybersecurity professional with more than 25 years of experience, specializing in red team tactics. As part of an authorized penetration test, and using your knowledge of OSINT and social engineering tactics, analyze the following sample job description for useful OSINT data that can be derived from it about the company such as systems used, programming languages used, job roles, locations, staff, etc. Be sure to include any correlations and conclusions you might draw. Only include data relevant to OSINT. Just provide me with the correlations and conclusions in your response.'},
        {"role": "user", "content": job_desc},
    ]

    # Call the OpenAI API
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )

    # Return the generated text
    return response.choices[0].message.content.strip()

# Check if the data directory exists
if not os.path.isdir('data'):
    print("Directory not found: 'data'")
    exit()

# Iterate through all files in data directory
try:
    for file_name in os.listdir('data'):
        with open('data/' + file_name, 'r') as file:
            job_desc = file.read()
            analysis = prompt_osint_analysis(job_desc)
            # Append the analysis to output.txt
            with open('output.txt', 'a') as output_file:
                output_file.write(analysis + '\n')
except FileNotFoundError:
    print("File not found.")
    exit()

# Check if the output file exists
if not os.path.isfile('output.txt'):
    print("File not found: 'output.txt'")
    exit()

# Read the output data
with open('output.txt', 'r') as file:
    output_data = file.read()

# Create a final report
def generate_final_report(output_data):
    # Define the conversation messages
    messages = [
        {"role": "system", "content": 'You are a cybersecurity professional with more than 25 years of experience, specializing in red team tactics. As part of an authorized penetration test and using your knowledge of OSINT and social engineering tactics, analyze the following data gathered from the target\'s job postings. Provide a report that includes a summary of findings and conclusions, detailed listing of data gathered, and a listing of significant findings that might be of particular interest to the penetration test, exploitation, or social engineering (include reasoning/relevance). Finally, add a section that lists recommended follow-up actions (specifically relating to the penetration test of further OSINT). Use markdown language formatting. Use the following report format: \n\n#OSINT Report Title\n\n##Summary\n\n##Details\n\n##Significant Findings\n\n##Recommended Follow-up Actions\n\nData:'},
        {"role": "user", "content": output_data},
    ]

    # Call the OpenAI API
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )

    # Return the generated text
    return response.choices[0].message.content.strip()

# Generate the report using the OpenAI API
report = generate_final_report(output_data)

# Save the report as a Word document
current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
docx_output_file = f"OSINT_Report_{current_datetime}.docx"

# Function to convert markdown text to a Word document
def markdown_to_docx(markdown_text: str, output_file: str):
    document = Document()

    # Iterate through the lines of the markdown text
    for line in markdown_text.split('\n'):
        # Add headings based on the markdown heading levels
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        # Add paragraphs for other text
        else:
            document.add_paragraph(line)

    # Save the Word document
    document.save(output_file)

# Handle exceptions during the report generation
try:
    markdown_to_docx(report, docx_output_file)
    print("Report generated successfully!")
except Exception as e:
    print(f"An error occurred during the report generation: {e}")