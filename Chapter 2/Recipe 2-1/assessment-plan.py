import openai
import os
from docx import Document
from tqdm import tqdm
import threading
import time
from datetime import datetime

# Set up the OpenAI API
openai.api_key = os.getenv("OPENAI_API_KEY")

current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
assessment_name = f"Vuln_Assessment_Plan_{current_datetime}"

def read_user_input_file(file_path: str) -> dict:
    user_data = {}
    with open(file_path, 'r') as file:
        for line in file:
            key, value = line.strip().split(':')
            user_data[key.strip()] = value.strip()
    return user_data

user_data_file = "assessment_data.txt" 
user_data = read_user_input_file(user_data_file)

# Function to generate a report using the OpenAI API
def generate_report(network_size, 
                    number_of_nodes, 
                    type_of_devices, 
                    special_devices, 
                    operating_systems, 
                    network_topology, 
                    access_controls, 
                    previous_security_incidents,
                    compliance_requirements,
                    business_critical_assets,
                    data_classification,
                    goals,
                    timeline,
                    team,
                    deliverables,
                    audience: str) -> str:

    # Define the conversation messages
    messages = [
        {"role": "system", "content": "You are a cybersecurity professional specializing in vulnerability assessment."},
        {"role": "user", "content": f'Using cybersecurity industry standards and best practices, create a complete and detailed assessment plan (not a penetration test) that includes: Introduction, outline of the process/methodology, tools needed, and a very detailed multi-layered outline of the steps. Provide a thorough and descriptive introduction and as much detail and description as possible throughout the plan. The plan should not only assessment of technical vulnerabilities on systems but also policies, procedures, and compliance. It should include the use of scanning tools as well as configuration review, staff interviews, and site walk-around. All recommendations should following industry standard best practices and methods. The plan should be a minimum of 1500 words.\n\
        Create the plan so that it is specific for the following details:\n\
        Network Size: {network_size}\n\
        Number of Nodes: {number_of_nodes}\n\
        Type of Devices: {type_of_devices}\n\
        Specific systems or devices that need to be excluded from the assessment: {special_devices}\n\
        Operating Systems: {operating_systems}\n\
        Network Topology: {network_topology}\n\
        Access Controls: {access_controls}\n\
        Previous Security Incidents: {previous_security_incidents}\n\
        Compliance Requirements: {compliance_requirements}\n\
        Business Critical Assets: {business_critical_assets}\n\
        Data Classification: {data_classification}\n\
        Goals and objectives of the vulnerability assessment: {goals}\n\
        Timeline for the vulnerability assessment: {timeline}\n\
        Team: {team}\n\
        Expected deliverables of the assessment: {deliverables}\n\
        Audience: {audience}\n\
        Provide the plan using the following format and observe the markdown language:\n\
        #Vulnerability Assessment Plan\n\
        ##Introduction\n\
        Thorough Introduction to the plan including the scope, reasons for doing it, goals and objectives, and summary of the plan\n\
        ##Process/Methodology\n\
        Description and Outline of the process/Methodology\n\
        ##Tools Required\n\
        List of required tools and applications, with their descriptions and reasons needed\n\
        ##Assessment Steps\n\
        Detailed, multi-layered outline of the assessment steps'}
    ]

    # Call the OpenAI API
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )

    # Return the generated text
    return response['choices'][0]['message']['content'].strip()

# Function to convert markdown text to a Word document
def markdown_to_docx(markdown_text: str, output_file: str):
    document = Document()

    # Iterate through the lines of the markdown text
    for line in markdown_text.split('\n'):

        # Add headings based on the markdown heading levels
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        # Add paragraphs for other text
        else:
            document.add_paragraph(line)

    # Save the Word document
    document.save(output_file)

# Function to display elapsed time while waiting for the API call
def display_elapsed_time():
    start_time = time.time()
    while not api_call_completed:
        elapsed_time = time.time() - start_time
        print(f"\rCommunicating with the API - Elapsed time: {elapsed_time:.2f} seconds", end="")
        time.sleep(1)

api_call_completed = False
elapsed_time_thread = threading.Thread(target=display_elapsed_time)
elapsed_time_thread.start()

# Handle exceptions during the API call
try:
    # Generate the report using the OpenAI API
    report = generate_report(
    user_data["Network Size"],
    user_data["Number of Nodes"],
    user_data["Type of Devices"],
    user_data["Specific systems or devices that need to be excluded from the assessment"],
    user_data["Operating Systems"],
    user_data["Network Topology"],
    user_data["Access Controls"],
    user_data["Previous Security Incidents"],
    user_data["Compliance Requirements"],
    user_data["Business Critical Assets"],
    user_data["Data Classification"],
    user_data["Goals and objectives of the vulnerability assessment"],
    user_data["Timeline for the vulnerability assessment"],
    user_data["Team"],
    user_data["Expected deliverables of the assessment"],
    user_data["Audience"]
    )

    api_call_completed = True
    elapsed_time_thread.join()
except Exception as e:
    api_call_completed = True
    elapsed_time_thread.join()
    print(f"\nAn error occurred during the API call: {e}")
    exit()

# Save the report as a Word document
docx_output_file = f"{assessment_name}_report.docx"

# Handle exceptions during the report generation
try:
    with tqdm(total=1, desc="Generating plan") as pbar:
        markdown_to_docx(report, docx_output_file)
        pbar.update(1)
    print("\nPlan generated successfully!")
except Exception as e:
    print(f"\nAn error occurred during the plan generation: {e}")
