import openai
from openai import OpenAI # New import required for the updated API call
import os
from docx import Document
from tqdm import tqdm
import threading
import time

# Set up the OpenAI API
openai.api_key = os.getenv("OPENAI_API_KEY")

# Function to generate a report using the OpenAI API
def generate_report(threat_name: str) -> str:

    # Define the conversation messages
    messages = [
        {"role": "system", "content": "You are a professional cyber threat analyst and MITRE ATT&CK Framework expert."},
        {"role": "user", "content": f'Provide a detailed report about {threat_name}, using the following template (and proper markdown language formatting, headings, bold keywords, tables, etc.):\n\n\
        Threat Name (Heading 1)\n\n\
        Summary (Heading 2)\n\
        Short executive summary\n\n\
        Details (Heading 2)\n\
        Description and details including history/background, discovery, characteristics and TTPs, known incidents\n\n\
        MITRE ATT&CK TTPs (Heading 2)\n\
        Table containing all of the known MITRE ATT&CK TTPs that the {threat_name} attack uses. Include the following collumns: Tactic, Technique ID, Technique Name, Procedure (How {threat_name} uses it)\n\n\
        Indicators of Compromise (Heading 2)\n\
        Table containing all of the known indicators of compromise. Include the following collumns: Type, Value, Description\n\n\  '}
    ]

    client = OpenAI() # New client initialization required for the updated API call

    # Call the OpenAI API
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )

    # Return the generated text
    return response.choices[0].message.content.strip()

# Function to convert markdown text to a Word document
def markdown_to_docx(markdown_text: str, output_file: str):
    document = Document()

    # Variables to keep track of the current table
    table = None
    in_table = False

    # Iterate through the lines of the markdown text
    for line in markdown_text.split('\n'):

        # Add headings based on the markdown heading levels
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        # Handle tables in the markdown text
        elif line.startswith('|'):
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table = document.add_table(rows=1, cols=len(row), style='Table Grid')
                for i, cell in enumerate(row):
                    table.cell(0, i).text = cell
            else:
                if len(row) != len(table.columns):  # If row length doesn't match table, it's a separator
                    continue
                new_row = table.add_row()
                for i, cell in enumerate(row):
                    new_row.cells[i].text = cell
        # Add paragraphs for other text
        else:
            if in_table:
                in_table = False
                table = None
            document.add_paragraph(line)

    # Save the Word document
    document.save(output_file)

# Function to extract tables from the markdown text
def extract_tables(markdown_text: str):
    tables = []
    current_table = []

    # Iterate through the lines of the markdown text
    for line in markdown_text.split('\n'):
        # Check if the line is part of a table
        if line.startswith('|'):
            current_table.append(line)
        # If the table ends, save it to the tables list
        elif current_table:
            tables.append('\n'.join(current_table))
            current_table = []

    return tables

# Function to display elapsed time while waiting for the API call
def display_elapsed_time():
    start_time = time.time()
    while not api_call_completed:
        elapsed_time = time.time() - start_time
        print(f"\rCommunicating with the API - Elapsed time: {elapsed_time:.2f} seconds", end="")
        time.sleep(1)

# Get user input
threat_name = input("Enter the name of a cyber threat: ")

api_call_completed = False
elapsed_time_thread = threading.Thread(target=display_elapsed_time)
elapsed_time_thread.start()

# Handle exceptions during the API call
try:
    # Generate the report using the OpenAI API
    report = generate_report(threat_name)
    api_call_completed = True
    elapsed_time_thread.join()
except Exception as e:
    api_call_completed = True
    elapsed_time_thread.join()
    print(f"\nAn error occurred during the API call: {e}")
    exit()

# Save the report as a Word document
docx_output_file = f"{threat_name}_report.docx"

# Handle exceptions during the report generation
try:
    with tqdm(total=1, desc="Generating report and files") as pbar:
        markdown_to_docx(report, docx_output_file)
    print("\nReport and tables generated successfully!")
except Exception as e:
    print(f"\nAn error occurred during the report generation: {e}")
