import openai
from openai import OpenAI # New import required for the updated API call
import os
from docx import Document

# Set up the OpenAI API
openai.api_key = os.getenv("OPENAI_API_KEY")

# Define the structure of the documents
design_doc_structure = [
    "Introduction",
    "Software Architecture",
    "Function Descriptions",
    "Flow Diagrams"
]

user_guide_structure = [
    "Introduction",
    "Installation Guide",
    "Usage Guide",
    "Troubleshooting"
]

def generate_section_content(section_title: str, source_code: str) -> str:
    print(f"Generating the {section_title} section content.\n")
    messages = [
        {"role": "system", "content": f"You are an experienced software engineer with extensive knowledge in writing {section_title} sections for design documents."},
        {"role": "user", "content": f"Please generate a {section_title} section for the following Python code:\n\n{source_code}"}
    ]

    client = OpenAI() # New client initialization required for the updated API call
    
    response = client.chat.completions.create( # Updated API call
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=2048,
        n=1,
        stop=None,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip() # Updated API call

def write_to_word_document(document: Document, title: str, content: str):
    document.add_heading(title, level=1)
    document.add_paragraph(content)

# Load the source code
with open('source_code_commented.py', 'r') as file:
    source_code = file.read()

# Create the design document
design_document = Document()

for section in design_doc_structure:
    section_content = generate_section_content(section, source_code)
    write_to_word_document(design_document, section, section_content)

design_document.save('DesignDocument.docx')

# Create the user guide
user_guide = Document()

for section in user_guide_structure:
    section_content = generate_section_content(section, source_code)
    write_to_word_document(user_guide, section, section_content)

print("The design document and user guide have been created.")
user_guide.save('UserGuide.docx')
